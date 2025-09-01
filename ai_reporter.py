# ai_reporter.py
import openai
import requests
import json
from docx import Document
from fpdf import FPDF
import sqlite3
import os
from database import db
from datetime import datetime
import time
import logging
import re
import html
import mdformat

# Setup logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class AIReporter:
    def __init__(self):
        self.api_key = db.get_api_key('openrouter')
        if self.api_key:
            openai.api_key = self.api_key
            openai.api_base = "https://openrouter.ai/api/v1"
        logger.info("AIReporter initialized")
    
    def set_api_key(self, api_key):
        """Set API key untuk OpenRouter"""
        self.api_key = api_key
        openai.api_key = api_key
        openai.api_base = "https://openrouter.ai/api/v1"
        db.save_api_key('openrouter', api_key)
        logger.info("API key set successfully")
    
    def get_available_models(self):
        """Dapatkan daftar model yang tersedia dari OpenRouter"""
        try:
            if not self.api_key:
                logger.warning("No API key found, returning default free models")
                # Return default free models jika tidak ada API key
                return self.get_default_free_models()
            
            headers = {
                "Authorization": f"Bearer {self.api_key}",
                "Content-Type": "application/json"
            }
            
            logger.info("Fetching available models from OpenRouter API")
            response = requests.get(
                "https://openrouter.ai/api/v1/models",
                headers=headers,
                timeout=15
            )
            
            if response.status_code == 200:
                data = response.json()
                models = data.get('data', [])
                logger.info(f"Retrieved {len(models)} models from API")
                
                # Filter dan format model
                formatted_models = []
                for model in models:
                    # Cek apakah model gratis atau memiliki rate limit yang baik
                    pricing = model.get('pricing', {})
                    is_free = (
                        pricing.get('prompt', '0') == '0' and 
                        pricing.get('completion', '0') == '0'
                    )
                    
                    model_info = {
                        'id': model.get('id'),
                        'name': model.get('name', model.get('id')),
                        'description': model.get('description', 'Tidak ada deskripsi'),
                        'context_length': model.get('context_length', 0),
                        'is_free': is_free,
                        'pricing': pricing
                    }
                    formatted_models.append(model_info)
                    logger.debug(f"Model: {model_info['id']} - Context: {model_info['context_length']}")
                
                # Urutkan: free models dulu, lalu berdasarkan context length
                formatted_models.sort(key=lambda x: (-x['is_free'], -x['context_length']))
                logger.info("Models sorted successfully")
                return formatted_models
            else:
                logger.error(f"API Error: {response.status_code} - {response.text}")
                return self.get_default_free_models()
                
        except requests.exceptions.Timeout:
            logger.error("Timeout saat mengambil daftar model")
            return self.get_default_free_models()
        except requests.exceptions.RequestException as e:
            logger.error(f"Request error: {e}")
            return self.get_default_free_models()
        except Exception as e:
            logger.error(f"Error getting models: {e}")
            return self.get_default_free_models()
    
    def get_default_free_models(self):
        """Return default free models jika API tidak tersedia"""
        logger.info("Returning default free models")
        return [
            {
                'id': 'mistralai/mistral-7b-instruct:free',
                'name': 'Mistral 7B Instruct (Free)',
                'description': 'Model dasar yang cepat dan gratis',
                'context_length': 32768,
                'is_free': True,
                'pricing': {'prompt': '0', 'completion': '0'}
            },
            {
                'id': 'google/gemma-2-9b-it:free',
                'name': 'Google Gemma 2 9B (Free)',
                'description': 'Model Google dengan kualitas baik',
                'context_length': 8192,
                'is_free': True,
                'pricing': {'prompt': '0', 'completion': '0'}
            },
            {
                'id': 'microsoft/phi-3-mini-128k-instruct:free',
                'name': 'Microsoft Phi-3 Mini (Free)',
                'description': 'Model Microsoft dengan context panjang',
                'context_length': 128000,
                'is_free': True,
                'pricing': {'prompt': '0', 'completion': '0'}
            },
            {
                'id': 'openchat/openchat-7b:free',
                'name': 'OpenChat 7B (Free)',
                'description': 'Model open source yang responsif',
                'context_length': 8192,
                'is_free': True,
                'pricing': {'prompt': '0', 'completion': '0'}
            }
        ]
    
    def get_user_models(self):
        """Dapatkan model yang pernah digunakan user"""
        logger.info("Getting user models from database")
        return db.get_user_models()
    
    def save_user_model(self, model_id, model_name=None):
        """Simpan model yang digunakan user"""
        logger.info(f"Saving user model: {model_id}")
        db.save_user_model(model_id, model_name)
    
    def get_model_context_length(self, model_id):
        """Dapatkan context length dari model"""
        try:
            # Coba dapatkan dari API jika ada API key
            if self.api_key:
                headers = {
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                }
                
                logger.info(f"Fetching context length for model: {model_id}")
                response = requests.get(
                    "https://openrouter.ai/api/v1/models",
                    headers=headers,
                    timeout=10
                )
                
                if response.status_code == 200:
                    data = response.json()
                    models = data.get('data', [])
                    for model in models:
                        if model.get('id') == model_id:
                            context_length = model.get('context_length', 4096)
                            logger.info(f"Context length for {model_id}: {context_length}")
                            return context_length
            
            # Jika tidak ada API key atau gagal, gunakan default values
            default_contexts = {
                'mistralai/mistral-7b-instruct:free': 32768,
                'google/gemma-2-9b-it:free': 8192,
                'microsoft/phi-3-mini-128k-instruct:free': 128000,
                'openchat/openchat-7b:free': 8192,
                'openai/gpt-3.5-turbo': 4096,
                'openai/gpt-4': 8192,
                'openai/gpt-4-turbo': 128000,
                'openai/gpt-4o': 128000,
                'meta-llama/llama-3-70b-instruct': 8192,
                'meta-llama/llama-3-8b-instruct': 8192
            }
            
            context_length = default_contexts.get(model_id, 4096)
            logger.info(f"Using default context length for {model_id}: {context_length}")
            return context_length
        except Exception as e:
            logger.error(f"Error getting context length: {e}")
            return 4096  # Default context length
    
    def estimate_token_count(self, text):
        """Estimasi jumlah token dari teks (1 token ≈ 4 karakter)"""
        token_count = len(text) // 4
        logger.debug(f"Estimated token count: {token_count} for text length: {len(text)}")
        return token_count
    
    def adjust_text_to_context(self, text, model_id, max_tokens_for_response=1000):
        """Sesuaikan teks agar sesuai dengan context length model"""
        context_length = self.get_model_context_length(model_id)
        logger.info(f"Adjusting text for model {model_id} with context length: {context_length}")
        
        # Estimasi token untuk prompt sistem dan instruksi (sekitar 300 token)
        system_prompt_tokens = 300
        logger.debug(f"System prompt tokens reserved: {system_prompt_tokens}")
        
        # Hitung token yang tersedia untuk transkripsi
        available_tokens = context_length - system_prompt_tokens - max_tokens_for_response
        logger.debug(f"Available tokens for transcription: {available_tokens}")
        
        # Estimasi token dari teks
        text_tokens = self.estimate_token_count(text)
        logger.info(f"Original text tokens: {text_tokens}")
        
        if text_tokens <= available_tokens:
            logger.info("Text fits within context, no adjustment needed")
            return text  # Tidak perlu dipotong
        
        # Jika terlalu panjang, potong teks
        # Hitung karakter maksimal yang diizinkan
        max_chars = available_tokens * 4
        logger.info(f"Text too long, max chars allowed: {max_chars}")
        
        if max_chars > 1000:  # Minimal 1000 karakter
            adjusted_text = text[:max_chars] + "... (teks dipotong untuk sesuai dengan limit context model)"
        else:
            adjusted_text = text[:1000] + "... (teks dipotong untuk sesuai dengan limit context model)"
        
        logger.info(f"Adjusted text length: {len(adjusted_text)}")
        return adjusted_text
    
    def get_model_max_completion_tokens(self, model_id):
        """
        Mendapatkan batas maksimum token completion untuk model tertentu.
        Mengambil informasi dari API OpenRouter atau menggunakan nilai default.
        """
        try:
            logger.info(f"Getting max completion tokens for model: {model_id}")
            
            # Jika ada API key, coba dapatkan dari API
            if self.api_key:
                headers = {
                    "Authorization": f"Bearer {self.api_key}",
                    "Content-Type": "application/json"
                }
                
                response = requests.get(
                    "https://openrouter.ai/api/v1/models",
                    headers=headers,
                    timeout=10
                )
                
                if response.status_code == 200:
                    data = response.json()
                    models = data.get('data', [])
                    
                    for model in models:
                        if model.get('id') == model_id:
                            # Dapatkan max_completion_tokens dari top_provider
                            top_provider = model.get('top_provider', {})
                            max_completion_tokens = top_provider.get('max_completion_tokens')
                            
                            if max_completion_tokens:
                                logger.info(f"Max completion tokens for {model_id}: {max_completion_tokens}")
                                return max_completion_tokens
                            else:
                                # Jika tidak ada info spesifik, gunakan default berdasarkan context_length
                                context_length = model.get('context_length', 4096)
                                # Gunakan 1/4 dari context length sebagai estimasi max completion
                                estimated_max = min(context_length // 4, 4000)  # Maksimal 4000
                                logger.info(f"No max completion info, using estimated max: {estimated_max}")
                                return estimated_max
            
            # Jika tidak ada API key atau gagal, gunakan default values
            default_max_tokens = {
                'mistralai/mistral-7b-instruct:free': 4000,
                'qwen/qwen3-235b-a22b:free': 8000,
                'google/gemini-flash-1.5-8b:free': 8000,
                'microsoft/phi-3-mini-128k-instruct:free': 4000,
                'openchat/openchat-7b:free': 4000,
                'google/gemma-2-9b-it:free': 4000,
                'meta-llama/llama-3.1-8b-instruct:free': 4000,
                'mistralai/mistral-nemo:free': 4000,
                'openai/gpt-3.5-turbo': 4096,
                'openai/gpt-4': 4096,
                'openai/gpt-4-turbo': 8192,
                'openai/gpt-4o': 16384,
                'anthropic/claude-3-haiku:free': 4096,
                'anthropic/claude-3-sonnet:free': 4096,
                'meta-llama/llama-3-70b-instruct': 4096,
                'meta-llama/llama-3-8b-instruct': 4096
            }
            
            max_tokens = default_max_tokens.get(model_id, 4000)
            logger.info(f"Using default max tokens for {model_id}: {max_tokens}")
            return max_tokens
            
        except Exception as e:
            logger.error(f"Error getting max completion tokens: {e}")
            # Return nilai aman default
            return 4000

    def calculate_optimal_max_tokens(self, model_id, report_type="summary"):
        """
        Menghitung max_tokens optimal berdasarkan tipe laporan dan model.
        
        Args:
            model_id (str): ID model yang digunakan
            report_type (str): Tipe laporan ("summary", "analysis", "custom")
        
        Returns:
            int: Max tokens yang direkomendasikan
        """
        max_model_tokens = self.get_model_max_completion_tokens(model_id)
        
        # Tentukan max_tokens berdasarkan tipe laporan
        if report_type == "summary":
            # Untuk ringkasan, gunakan 50-70% dari kapasitas model
            recommended = min(int(max_model_tokens * 1), 5000)
        elif report_type == "analysis":
            # Untuk analisis, gunakan 60-80% dari kapasitas model
            recommended = min(int(max_model_tokens * 1), 6000)
        elif report_type == "custom":
            # Untuk laporan kustom, gunakan 70-90% dari kapasitas model
            recommended = min(int(max_model_tokens * 1), 8000)
        else:
            # Default untuk tipe lainnya
            recommended = min(int(max_model_tokens * 1), 4000)
        
        logger.info(f"Calculated optimal max_tokens for {report_type} with {model_id}: {recommended}")
        return recommended

    def format_report_content_for_document(self, content: str):
        """
        Memformat konten laporan untuk dokumen DOCX/PDF dengan menangani bagian <think>.
        
        Args:
            content (str): Konten laporan yang akan diformat
            
        Returns:
            tuple: (plain_content, html_content) - Konten yang telah diformat
        """
        logger.info("Formatting report content for document")
        
        # Normalisasi line endings
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # Ekstrak dan proses bagian <think>
        think_sections = []
        def extract_think_section(match):
            think_content = match.group(1)
            think_sections.append(think_content)
            # Return placeholder yang akan diganti nanti
            return f"<<<THINK_PLACEHOLDER_{len(think_sections)-1}>>>"
        
        # Ekstrak bagian <think>
        content = re.sub(r'<think>(.*?)</think>', extract_think_section, content, flags=re.DOTALL)
        
        # Bersihkan karakter yang tidak diinginkan tapi pertahankan struktur
        content = re.sub(r'[^\S\n]+$', '', content, flags=re.MULTILINE)  # hapus spasi di akhir baris
        content = re.sub(r'\n{3,}', '\n\n', content)  # maksimal 2 baris kosong berturut-turut
        
        # Konversi heading markdown ke format yang konsisten
        # H1 sampai H6
        for i in range(6, 0, -1):
            pattern = r'^' + ('#' * i) + r'\s+(.+)$'
            replacement = '#' * i + r' \1'
            content = re.sub(pattern, replacement, content, flags=re.MULTILINE)
        
        # Konversi bold dan italic
        content = re.sub(r'\*\*(.*?)\*\*', r'**\1**', content)  # bold
        content = re.sub(r'\*(.*?)\*', r'*\1*', content)        # italic
        
        # Konversi list
        content = re.sub(r'^\s*-\s+(.+)', r'- \1', content, flags=re.MULTILINE)  # unordered list
        content = re.sub(r'^\s*\d+\.\s+(.+)', r'1. \1', content, flags=re.MULTILINE)  # ordered list
        
        # Persiapkan versi HTML untuk dokumen yang mendukung HTML
        html_content = content
        
        # Konversi untuk HTML (jika diperlukan)
        # Heading
        for i in range(6, 0, -1):
            pattern = r'^' + ('#' * i) + r'\s+(.+)$'
            replacement = f'<h{i}>\\1</h{i}>'
            html_content = re.sub(pattern, replacement, html_content, flags=re.MULTILINE)
        
        # Bold dan italic
        html_content = re.sub(r'\*\*(.*?)\*\*', r'<strong>\1</strong>', html_content)
        html_content = re.sub(r'\*(.*?)\*', r'<em>\1</em>', html_content)
        
        # List
        html_content = re.sub(r'^-\s+(.+)', r'<li>\1</li>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'(<li>.*?</li>\s*)+', r'<ul>\n\g<0></ul>\n', html_content, flags=re.DOTALL)
        
        html_content = re.sub(r'^\d+\.\s+(.+)', r'<li>\1</li>', html_content, flags=re.MULTILINE)
        html_content = re.sub(r'(<li>.*?</li>\s*)+', r'<ol>\n\g<0></ol>\n', html_content, flags=re.DOTALL)
        
        # Masukkan kembali bagian think dengan format yang sesuai
        for i, think_content in enumerate(think_sections):
            # Format untuk konten teks biasa
            think_placeholder = f"<<<THINK_PLACEHOLDER_{i}>>>"
            formatted_think = f"\n[Proses Berpikir]\n{think_content}\n"
            content = content.replace(think_placeholder, formatted_think)
            
            # Format untuk HTML
            html_think = f'<div class="think-box"><h3>💭 Proses Berpikir</h3><p>{html.escape(think_content)}</p></div>'
            html_content = html_content.replace(think_placeholder, html_think)
        
        # Pastikan ada baris baru di akhir
        if not content.endswith('\n'):
            content += '\n'
        
        if not html_content.endswith('\n'):
            html_content += '\n'
        
        return content, html_content

    def clean_for_docx(self, content: str) -> str:
        """
        Membersihkan dan memformat konten untuk format DOCX menggunakan mdformat.
        """
        try:
            # Gunakan mdformat untuk memformat konten Markdown
            formatted_content = mdformat.text(content)
            
            # Lanjutkan dengan pembersihan karakter khusus untuk DOCX
            replacements = {
                '\u2013': '-',  # en dash
                '\u2014': '-',  # em dash
                '\u2018': "'",  # left single quotation mark
                '\u2019': "'",  # right single quotation mark
                '\u201c': '"',  # left double quotation mark
                '\u201d': '"',  # right double quotation mark
                '\u2026': '...', # horizontal ellipsis
            }
            
            for unicode_char, replacement in replacements.items():
                formatted_content = formatted_content.replace(unicode_char, replacement)
            
            # Hapus karakter kontrol yang tidak diizinkan di XML/DOCX
            cleaned_content = ""
            for char in formatted_content:
                if ord(char) == 0x09 or ord(char) == 0x0A or ord(char) == 0x0D or \
                   (0x20 <= ord(char) <= 0xD7FF) or \
                   (0xE000 <= ord(char) <= 0xFFFD) or \
                   (0x10000 <= ord(char) <= 0x10FFFF):
                    cleaned_content += char
                else:
                    # Ganti karakter tidak valid dengan spasi
                    cleaned_content += ' '
            
            return cleaned_content.strip()
        except Exception as e:
            # Jika mdformat gagal, fallback ke pembersihan dasar
            logger.warning(f"mdformat failed: {e}. Using basic cleaning.")
            return self.basic_clean_for_docx(content)

    def clean_for_pdf(self, content: str) -> str:
        """
        Membersihkan dan memformat konten untuk format PDF menggunakan mdformat.
        """
        try:
            # Gunakan mdformat untuk memformat konten Markdown
            formatted_content = mdformat.text(content)
            
            # Untuk PDF, kita bisa mempertahankan lebih banyak karakter Unicode
            # Tapi tetap perlu membersihkan karakter kontrol
            cleaned_content = ""
            for char in formatted_content:
                if ord(char) >= 32 or char in '\n\r\t':
                    cleaned_content += char
                else:
                    # Ganti karakter kontrol dengan spasi
                    cleaned_content += ' '
            
            return cleaned_content.strip()
        except Exception as e:
            # Jika mdformat gagal, fallback ke pembersihan dasar
            logger.warning(f"mdformat failed: {e}. Using basic cleaning.")
            return self.basic_clean_for_pdf(content)

    # Fungsi pembersihan dasar sebagai fallback (contoh untuk DOCX)
    def basic_clean_for_docx(self, content: str) -> str:
        # Normalisasi line ending
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        # Tambahkan pembersihan karakter Unicode dan kontrol seperti sebelumnya
        replacements = {
            '\u2013': '-',  # en dash
            '\u2014': '-',  # em dash
            # ... tambahkan lainnya
        }
        for unicode_char, replacement in replacements.items():
            content = content.replace(unicode_char, replacement)
        
        cleaned_content = ""
        for char in content:
            if ord(char) == 0x09 or ord(char) == 0x0A or ord(char) == 0x0D or \
               (0x20 <= ord(char) <= 0xD7FF) or \
               (0xE000 <= ord(char) <= 0xFFFD) or \
               (0x10000 <= ord(char) <= 0x10FFFF):
                cleaned_content += char
            else:
                cleaned_content += ' '
        return cleaned_content.strip()

    # Fungsi pembersihan dasar sebagai fallback (contoh untuk PDF)
    def basic_clean_for_pdf(self, content: str) -> str:
        # Normalisasi line ending
        content = content.replace('\r\n', '\n').replace('\r', '\n')
        
        cleaned_content = ""
        for char in content:
            if ord(char) >= 32 or char in '\n\r\t':
                cleaned_content += char
            else:
                cleaned_content += ' '
        return cleaned_content.strip()

    def generate_summary(self, transcription_text, model_id="mistralai/mistral-7b-instruct:free"):
        """Generate ringkasan dari transkripsi"""
        logger.info(f"Generating summary with model: {model_id}")
        
        if not self.api_key:
            logger.error("API key not found")
            raise Exception("API key tidak ditemukan. Silakan set API key terlebih dahulu.")
        
        # Simpan model yang digunakan
        self.save_user_model(model_id, model_id)
        
        # Sesuaikan teks dengan context length model
        adjusted_text = self.adjust_text_to_context(transcription_text, model_id, 1000)
        logger.info(f"Adjusted text length: {len(adjusted_text)}")
        
        prompt = f"""
        Buat ringkasan yang komprehensif dari transkripsi berikut dalam bahasa Indonesia:
        
        Transkripsi:
        {adjusted_text}
        
        Harap berikan ringkasan yang mencakup:
        1. Poin-poin utama yang dibahas
        2. Kesimpulan penting
        3. Rekomendasi jika ada
        
        Ringkasan harus jelas, terstruktur, dan informatif.
        """
        
        logger.info(f"Prompt length: {len(prompt)} characters")
        
        try:
            # Validasi model ID
            if not model_id or model_id == "":
                model_id = "mistralai/mistral-7b-instruct:free"
                logger.warning("Using default model due to empty model_id")
            
            # Hitung max_tokens optimal
            max_tokens = self.calculate_optimal_max_tokens(model_id, "summary")
            logger.info(f"Using max_tokens: {max_tokens}")
            
            logger.info("Sending request to OpenAI API")
            response = openai.ChatCompletion.create(
                model=model_id,
                messages=[
                    {"role": "system", "content": "Anda adalah asisten yang ahli dalam membuat ringkasan dari transkripsi dalam bahasa Indonesia."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,  # Gunakan nilai yang dihitung
                temperature=0.7,
                timeout=120  # 2 menit timeout
            )
            
            result = response.choices[0].message.content.strip()
            logger.info(f"Summary generated successfully. Response length: {len(result)} characters")
            return result
        except openai.error.Timeout as e:
            logger.error(f"Timeout saat menghasilkan ringkasan: {str(e)}")
            raise Exception(f"Timeout saat menghasilkan ringkasan: {str(e)}")
        except openai.error.APIError as e:
            logger.error(f"API Error: {str(e)}")
            raise Exception(f"API Error: {str(e)}")
        except openai.error.RateLimitError as e:
            logger.error(f"Rate limit exceeded: {str(e)}")
            raise Exception(f"Rate limit exceeded: {str(e)}")
        except Exception as e:
            logger.error(f"Gagal menghasilkan ringkasan: {str(e)}")
            raise Exception(f"Gagal menghasilkan ringkasan: {str(e)}")

    def generate_analysis(self, transcription_text, analysis_type="general", model_id="mistralai/mistral-7b-instruct:free"):
        """Generate analisis dari transkripsi"""
        logger.info(f"Generating analysis with model: {model_id}, type: {analysis_type}")
        
        if not self.api_key:
            logger.error("API key not found")
            raise Exception("API key tidak ditemukan.")
        
        # Simpan model yang digunakan
        self.save_user_model(model_id, model_id)
        
        # Sesuaikan teks dengan context length model
        adjusted_text = self.adjust_text_to_context(transcription_text, model_id, 1500)
        logger.info(f"Adjusted text length: {len(adjusted_text)}")
        
        analysis_prompts = {
            "general": """
            Lakukan analisis komprehensif terhadap transkripsi berikut dalam bahasa Indonesia:
            
            Transkripsi:
            {text}
            
            Berikan analisis yang mencakup:
            1. Tema utama yang dibahas
            2. Sentimen keseluruhan
            3. Poin-poin penting
            4. Insight menarik
            5. Rekomendasi berdasarkan analisis
            """,
            
            "sentiment": """
            Analisis sentimen dari transkripsi berikut dalam bahasa Indonesia:
            
            Transkripsi:
            {text}
            
            Berikan analisis yang mencakup:
            1. Sentimen keseluruhan (positif/negatif/netral)
            2. Emosi yang dominan
            3. Poin-poin dengan sentimen berbeda
            4. Rekomendasi untuk meningkatkan komunikasi
            """,
            
            "keypoints": """
            Identifikasi poin-poin kunci dari transkripsi berikut dalam bahasa Indonesia:
            
            Transkripsi:
            {text}
            
            Ekstrak dan organisir poin-poin penting yang mencakup:
            1. Keputusan penting yang diambil
            2. Isu-isu yang dibahas
            3. Tindakan yang disepakati
            4. Waktu dan tanggal penting
            5. Pihak-pihak yang terlibat
            """
        }
        
        prompt_template = analysis_prompts.get(analysis_type, analysis_prompts["general"])
        prompt = prompt_template.format(text=adjusted_text)
        logger.info(f"Prompt length: {len(prompt)} characters")
        
        try:
            # Validasi model ID
            if not model_id or model_id == "":
                model_id = "mistralai/mistral-7b-instruct:free"
                logger.warning("Using default model due to empty model_id")
            
            # Hitung max_tokens optimal
            max_tokens = self.calculate_optimal_max_tokens(model_id, "analysis")
            logger.info(f"Using max_tokens: {max_tokens}")
            
            logger.info("Sending request to OpenAI API")
            response = openai.ChatCompletion.create(
                model=model_id,
                messages=[
                    {"role": "system", "content": "Anda adalah analis yang ahli dalam menganalisis transkripsi dalam bahasa Indonesia."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=max_tokens,  # Gunakan nilai yang dihitung
                temperature=0.7,
                timeout=120
            )
            
            result = response.choices[0].message.content.strip()
            logger.info(f"Analysis generated successfully. Response length: {len(result)} characters")
            return result
        except openai.error.Timeout as e:
            logger.error(f"Timeout saat menghasilkan analisis: {str(e)}")
            raise Exception(f"Timeout saat menghasilkan analisis: {str(e)}")
        except openai.error.APIError as e:
            logger.error(f"API Error: {str(e)}")
            raise Exception(f"API Error: {str(e)}")
        except Exception as e:
            logger.error(f"Gagal menghasilkan analisis: {str(e)}")
            raise Exception(f"Gagal menghasilkan analisis: {str(e)}")

    def generate_custom_report(self, transcription_text, custom_prompt, model_id="mistralai/mistral-7b-instruct:free"):
        """Generate laporan kustom berdasarkan prompt user"""
        logger.info(f"Generating custom report with model: {model_id}")
        
        if not self.api_key:
            logger.error("API key not found")
            raise Exception("API key tidak ditemukan.")
        
        # Validasi input
        if not custom_prompt or not custom_prompt.strip():
            logger.error("Custom prompt is empty")
            raise Exception("Prompt kustom tidak boleh kosong.")
        
        # Simpan model yang digunakan
        self.save_user_model(model_id, model_id)
        
        # Sesuaikan teks dengan context length model
        adjusted_text = self.adjust_text_to_context(transcription_text, model_id, 2000)
        logger.info(f"Adjusted text length: {len(adjusted_text)}")
        
        full_prompt = f"""
        Berdasarkan transkripsi berikut dalam bahasa Indonesia:
        
        Transkripsi:
        {adjusted_text}
        
        Petunjuk pengguna:
        {custom_prompt}
        
        Harap berikan jawaban yang komprehensif dan terstruktur dalam bahasa Indonesia.
        Pastikan jawaban Anda relevan dengan transkripsi dan petunjuk yang diberikan.
        """
        
        logger.info(f"Full prompt length: {len(full_prompt)} characters")
        
        try:
            # Validasi model ID
            if not model_id or model_id == "":
                model_id = "mistralai/mistral-7b-instruct:free"
                logger.warning("Using default model due to empty model_id")
            
            # Hitung max_tokens optimal
            max_tokens = self.calculate_optimal_max_tokens(model_id, "custom")
            logger.info(f"Using max_tokens: {max_tokens}")
            
            logger.info("Sending request to OpenAI API")
            response = openai.ChatCompletion.create(
                model=model_id,
                messages=[
                    {"role": "system", "content": "Anda adalah asisten yang ahli dalam membuat laporan berdasarkan instruksi spesifik dalam bahasa Indonesia."},
                    {"role": "user", "content": full_prompt}
                ],
                max_tokens=max_tokens,  # Gunakan nilai yang dihitung
                temperature=0.7,
                timeout=180  # 3 menit timeout untuk custom report
            )
            
            result = response.choices[0].message.content.strip()
            logger.info(f"Custom report generated successfully. Response length: {len(result)} characters")
            return result
        except openai.error.Timeout as e:
            logger.error(f"Timeout saat menghasilkan laporan kustom: {str(e)}")
            raise Exception(f"Timeout saat menghasilkan laporan kustom: {str(e)}")
        except openai.error.APIError as e:
            logger.error(f"API Error: {str(e)}")
            raise Exception(f"API Error: {str(e)}")
        except Exception as e:
            logger.error(f"Gagal menghasilkan laporan kustom: {str(e)}")
            raise Exception(f"Gagal menghasilkan laporan kustom: {str(e)}")
    
    def create_docx_report(self, title, content, output_path):
        """Buat laporan dalam format DOCX"""
        logger.info(f"Creating DOCX report: {output_path}")
        try:
            # Bersihkan konten terlebih dahulu
            cleaned_content = self.clean_for_docx(content)
            
            doc = Document()
            
            # Judul
            doc.add_heading(title, 0)
            
            # Metadata
            doc.add_paragraph(f'Dibuat pada: {datetime.now().strftime("%d %B %Y, %H:%M:%S")}')
            doc.add_paragraph('')
            
            # Konten - split berdasarkan baris baru dan proses markdown
            lines = cleaned_content.split('\n')
            in_list = False
            list_items = []
            
            for line in lines:
                line = line.strip()
                
                # Skip empty lines
                if not line:
                    # Jika dalam list, tambahkan list items
                    if in_list and list_items:
                        paragraph = doc.add_paragraph()
                        for item in list_items:
                            paragraph.add_run('• ').bold = True
                            paragraph.add_run(item + '\n')
                        list_items = []
                        in_list = False
                    continue
                
                # Cek heading
                if line.startswith('# '):
                    if in_list and list_items:
                        paragraph = doc.add_paragraph()
                        for item in list_items:
                            paragraph.add_run('• ').bold = True
                            paragraph.add_run(item + '\n')
                        list_items = []
                        in_list = False
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    if in_list and list_items:
                        paragraph = doc.add_paragraph()
                        for item in list_items:
                            paragraph.add_run('• ').bold = True
                            paragraph.add_run(item + '\n')
                        list_items = []
                        in_list = False
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    if in_list and list_items:
                        paragraph = doc.add_paragraph()
                        for item in list_items:
                            paragraph.add_run('• ').bold = True
                            paragraph.add_run(item + '\n')
                        list_items = []
                        in_list = False
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    # List item
                    in_list = True
                    list_items.append(line[2:])
                else:
                    # Paragraf biasa
                    if in_list and list_items:
                        paragraph = doc.add_paragraph()
                        for item in list_items:
                            paragraph.add_run('• ').bold = True
                            paragraph.add_run(item + '\n')
                        list_items = []
                        in_list = False
                    doc.add_paragraph(line)
            
            # Tambahkan list items yang tersisa
            if in_list and list_items:
                paragraph = doc.add_paragraph()
                for item in list_items:
                    paragraph.add_run('• ').bold = True
                    paragraph.add_run(item + '\n')
            
            doc.save(output_path)
            logger.info(f"DOCX report saved successfully: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Gagal membuat file DOCX: {str(e)}")
            raise Exception(f"Gagal membuat file DOCX: {str(e)}")
    
    def create_pdf_report(self, title, content, output_path):
        """Buat laporan dalam format PDF"""
        logger.info(f"Creating PDF report: {output_path}")
        try:
            # Bersihkan konten terlebih dahulu
            cleaned_content = self.clean_for_pdf(content)
            
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Title
            pdf.set_font('Arial', 'B', 16)
            try:
                pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
            except:
                pdf.cell(0, 10, title, ln=True, align='C')
            pdf.ln(10)
            
            # Metadata
            pdf.set_font('Arial', '', 10)
            try:
                pdf.cell(0, 10, f'Dibuat pada: {datetime.now().strftime("%d %B %Y, %H:%M:%S")}'.encode('latin-1', 'replace').decode('latin-1'), ln=True)
            except:
                pdf.cell(0, 10, f'Dibuat pada: {datetime.now().strftime("%d %B %Y, %H:%M:%S")}', ln=True)
            pdf.ln(10)
            
            # Content
            pdf.set_font('Arial', '', 12)
            
            # Split content and add to PDF
            lines = cleaned_content.split('\n')
            for line in lines:
                # Handle encoding issues
                try:
                    clean_line = line.encode('latin-1', 'replace').decode('latin-1')
                    # Handle long lines
                    if len(clean_line) > 80:
                        # Split into multiple lines
                        while len(clean_line) > 80:
                            pdf.cell(0, 10, clean_line[:80], ln=True)
                            clean_line = clean_line[80:]
                        if clean_line:
                            pdf.cell(0, 10, clean_line, ln=True)
                    else:
                        pdf.cell(0, 10, clean_line, ln=True)
                except:
                    try:
                        pdf.cell(0, 10, "Error encoding line", ln=True)
                    except:
                        pass
            
            pdf.output(output_path)
            logger.info(f"PDF report saved successfully: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Gagal membuat file PDF: {str(e)}")
            raise Exception(f"Gagal membuat file PDF: {str(e)}")

# Inisialisasi AI Reporter
ai_reporter = AIReporter()